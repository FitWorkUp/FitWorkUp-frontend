from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\ronal\Downloads\Projetos\FitWorkUp-frontend")
MEDIA = ROOT / ".codex-doc-review" / "expanded" / "media"
LOGO = ROOT / ".codex-doc-review" / "expanded" / "word" / "media" / "image1.png"
OUTPUT = ROOT / "Relatório de Avaliação de Usabilidade - Revisado.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_body(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if bold_label and text.startswith(bold_label):
        p.add_run(bold_label).bold = True
        p.add_run(text[len(bold_label):])
    else:
        p.add_run(text)
    return p


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.add_run(label).bold = True
    p.add_run(text)
    return p


def add_figure(doc, path, number, title, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Cm(7.2))
    set_alt_text(shape, alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run(f"Figura {number} – {title}")
    r.bold = True
    r.font.size = Pt(10)
    src = doc.add_paragraph("Fonte: Elaborado pelo autor (2026).")
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(8)
    src.runs[0].font.size = Pt(9)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)

for style_name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 12)):
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)

styles["Heading 1"].paragraph_format.space_before = Pt(12)
styles["Heading 1"].paragraph_format.space_after = Pt(8)
styles["Heading 1"].paragraph_format.keep_with_next = True
styles["Heading 2"].paragraph_format.space_before = Pt(10)
styles["Heading 2"].paragraph_format.space_after = Pt(6)
styles["Heading 2"].paragraph_format.keep_with_next = True

# Capa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = p.add_run().add_picture(str(LOGO), width=Cm(3.0))
set_alt_text(shape, "Logomarca do Instituto Federal da Bahia")

p = doc.add_paragraph("INSTITUTO FEDERAL DE EDUCAÇÃO, CIÊNCIA E TECNOLOGIA DA BAHIA\nCAMPUS SANTO ANTÔNIO DE JESUS")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph("CURSO SUPERIOR DE TECNOLOGIA EM ANÁLISE E DESENVOLVIMENTO DE SISTEMAS")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

for _ in range(1):
    doc.add_paragraph()
p = doc.add_paragraph("RONALDO CORREIA COUTO")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph("RELATÓRIO DE AVALIAÇÃO DE USABILIDADE E ACESSIBILIDADE:\nESTUDO DE CASO DO APLICATIVO FITWORKUP")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(14)

for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph("SANTO ANTÔNIO DE JESUS – BA\n2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

# Folha de rosto
doc.add_section(WD_SECTION.NEW_PAGE)
p = doc.add_paragraph("RONALDO CORREIA COUTO")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph("RELATÓRIO DE AVALIAÇÃO DE USABILIDADE E ACESSIBILIDADE:\nESTUDO DE CASO DO APLICATIVO FITWORKUP")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(14)
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph("Relatório acadêmico apresentado como requisito de avaliação da disciplina Interface Homem-Computador, ministrada pelo professor João Manoel Andrade Moreira.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.line_spacing = 1.0
for _ in range(7):
    doc.add_paragraph()
p = doc.add_paragraph("SANTO ANTÔNIO DE JESUS – BA\n2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

body_section = doc.add_section(WD_SECTION.NEW_PAGE)

doc.add_heading("1 INTRODUÇÃO E OBJETO DE ESTUDO", level=1)
add_body(doc, "Este relatório apresenta uma avaliação preliminar da interface do FitWorkUp, aplicativo Android voltado ao registro e ao acompanhamento de caminhadas e corridas. O sistema reúne metas, histórico, ranking semanal, recompensas virtuais, loja e perfil social. A análise concentra-se na forma como essas funções são comunicadas e operadas pelo usuário, e não na eficácia clínica ou comportamental do produto.")
add_body(doc, "O objetivo é identificar problemas de usabilidade e acessibilidade observáveis no protótipo e propor melhorias compatíveis com o estágio atual do desenvolvimento. As críticas articulam os temas de ética, ergonomia, cultura, psicologia, cor, acessibilidade, iconografia, tipografia e modelos de avaliação de interface. Como a inspeção foi realizada sem participantes, seus resultados constituem um diagnóstico técnico inicial e deverão ser confirmados posteriormente por testes com usuários representativos.")
add_body(doc, "O recorte considera as telas inicial, ranking, loja e perfil exibidas nas capturas do protótipo Android em orientação vertical. Quando necessário, a inspeção visual foi confrontada com a implementação atual em Jetpack Compose para evitar conclusões baseadas apenas na aparência das imagens.")

doc.add_heading("2 METODOLOGIA", level=1)
doc.add_heading("2.1 Tipo e limites da avaliação", level=2)
add_body(doc, "Foi adotada uma inspeção heurística qualitativa conduzida pelo autor. Não houve recrutamento de participantes, aplicação de questionários, cronometragem de tarefas ou coleta de métricas comportamentais. Portanto, o relatório não apresenta resultados de teste de usabilidade; ele formula hipóteses de problemas que devem orientar correções e uma avaliação empírica posterior.")

doc.add_heading("2.2 Materiais e procedimento", level=2)
add_body(doc, "A inspeção utilizou quatro capturas representativas do protótipo e a leitura dos componentes de interface correspondentes. O procedimento consistiu em: identificar a tarefa principal de cada tela; verificar clareza, consistência, legibilidade, área de toque, adaptação do conteúdo e comunicação de estados; relacionar cada ocorrência a princípios reconhecidos; e registrar uma recomendação verificável.")

doc.add_heading("2.3 Referenciais e classificação", level=2)
add_body(doc, "Foram utilizadas as heurísticas de Nielsen, as diretrizes WCAG 2.2, recomendações de acessibilidade do Android e conceitos de carga cognitiva, comparação social, design enganoso e lei de Fitts. A severidade segue a escala de Nielsen: 0 (não é problema), 1 (cosmético), 2 (pequeno), 3 (grave) e 4 (catastrófico). A classificação considera frequência estimada, impacto e persistência; por ter sido atribuída por um único avaliador, deve ser tratada como prioridade preliminar, não como medida definitiva.")

doc.add_heading("3 INTERFACES AVALIADAS", level=1)
add_figure(doc, MEDIA / "image.jpg", 1, "Tela inicial do FitWorkUp", "Captura da tela inicial com saudação, meta de passos, resumo diário e calendário mensal")
add_figure(doc, MEDIA / "image2.jpg", 2, "Ranking semanal", "Captura da classificação semanal com pódio e lista de participantes")
add_figure(doc, MEDIA / "image3.jpg", 3, "Loja virtual", "Captura da loja com itens cosméticos, multiplicadores e botões de ação")
add_figure(doc, MEDIA / "image4.jpg", 4, "Perfil do usuário", "Captura do perfil com avatar, nível, estatísticas, conquistas e conexões")

doc.add_heading("4 DIAGNÓSTICO E ANÁLISE CRÍTICA", level=1)

doc.add_heading("4.1 Crítica 1 – Legibilidade, contraste e comunicação acessível", level=2)
add_labeled(doc, "Problema observado: ", "no ranking, informações secundárias são apresentadas em tamanhos reduzidos, entre aproximadamente 9 sp e 11 sp, e algumas usam cor com transparência sobre fundos claros. Essa combinação enfraquece a leitura de passos, posição e demais dados, sobretudo para pessoas com baixa visão ou que aumentam o tamanho da fonte do sistema.")
add_labeled(doc, "Evidência: ", "a captura do ranking mostra textos pequenos no pódio e na lista. Na implementação, alguns textos secundários utilizam opacidade de 72%. Considerando as cores atuais, a combinação se aproxima de 3:1 em fundo claro, abaixo da relação de 4,5:1 indicada pela WCAG 2.2 para texto comum. Esse cálculo é uma verificação de projeto e deve ser confirmado nas cores efetivamente renderizadas em ambos os temas.")
add_labeled(doc, "Fundamentação: ", "o critério 1.4.3 da WCAG 2.2 estabelece contraste mínimo de 4,5:1 para texto comum e 3:1 para texto grande. O critério 1.4.1 também recomenda que a cor não seja o único meio de transmitir informação. As orientações de acessibilidade do Android reforçam a necessidade de conteúdo legível, semântica adequada e compatibilidade com tecnologias assistivas.")
add_labeled(doc, "Recomendação: ", "adotar tamanho mínimo equivalente ao estilo bodySmall para dados auxiliares; evitar reduzir contraste apenas por transparência; medir o contraste nos temas claro e escuro; manter texto, ícone e forma como sinais redundantes; verificar a interface com ampliação de fonte, Accessibility Scanner e TalkBack. Componentes personalizados de progresso devem expor descrição e estado quando essa informação não estiver disponível como texto próximo.")
add_labeled(doc, "Severidade: ", "3 – problema grave, porque afeta conteúdo recorrente e pode impedir a leitura por parte do público.")

doc.add_heading("4.2 Crítica 2 – Transparência ética e cultura de competição", level=2)
add_labeled(doc, "Problema observado: ", "o ranking apresenta perfis simulados, como bot_atlas, bot_flash, bot_ninja e bot_runner, na mesma estrutura visual dos usuários reais, sem um selo explícito de conteúdo simulado. Ainda que os nomes deem uma pista, o usuário não recebe uma explicação sobre a origem desses perfis nem sobre sua participação na classificação.")
add_labeled(doc, "Evidência: ", "a captura exibe contas automatizadas no pódio e na lista, com aparência equivalente às demais. Essa apresentação pode produzir uma impressão artificial de atividade social e de competição. Além disso, um ranking único privilegia a comparação pública e pode ter efeitos diferentes sobre iniciantes, pessoas com limitações físicas ou usuários menos competitivos.")
add_labeled(doc, "Fundamentação: ", "Gray et al. descrevem como decisões de interface podem induzir interpretações ou escolhas que favorecem o sistema em detrimento da compreensão do usuário. A teoria da comparação social de Festinger explica que avaliações pessoais são influenciadas pela comparação com outras pessoas, o que torna a apresentação do ranking relevante para a motivação e para a cultura criada pelo aplicativo.")
add_labeled(doc, "Recomendação: ", "identificar perfis simulados com selo “Bot” ou “Simulado”; exibir uma explicação breve sobre seu uso; impedir que bots ocupem posições do ranking oficial; e oferecer alternativas como ranking entre amigos, progresso pessoal e metas cooperativas. As regras de pontuação e reinício semanal também devem estar acessíveis na própria tela.")
add_labeled(doc, "Severidade: ", "2 – problema pequeno a moderado, pois não bloqueia a tarefa, mas pode reduzir confiança e adequação cultural.")

doc.add_heading("4.3 Crítica 3 – Ergonomia das áreas de toque", level=2)
add_labeled(doc, "Problema observado: ", "há controles compactos em cartões densos, incluindo botões da loja definidos com 44 dp de altura, valor inferior ao alvo mínimo de 48 dp recomendado para elementos interativos no Android. O risco cresce em um aplicativo de atividade física, que pode ser consultado em movimento e com menor precisão motora.")
add_labeled(doc, "Evidência: ", "na tela da loja, os botões “Resgatar” e “Ativar” ficam próximos às bordas dos cartões e a outros elementos. A implementação do botão contornado da loja fixa sua altura em 44 dp. Em contrapartida, componentes Material como IconButton podem ampliar automaticamente a área de toque; por isso, a crítica se limita aos controles cuja dimensão reduzida foi confirmada.")
add_labeled(doc, "Fundamentação: ", "a documentação de acessibilidade do Android recomenda alvos de toque de pelo menos 48 dp. Pela lei de Fitts, alvos menores e mais distantes exigem maior precisão e tendem a aumentar o tempo e a probabilidade de erro durante a seleção.")
add_labeled(doc, "Recomendação: ", "substituir altura fixa por heightIn(min = 48.dp), assegurar espaçamento entre ações, reservar margens internas adequadas e testar o uso com uma mão. A validação deve incluir inspeção automática e tarefas reais de compra, ativação e navegação executadas em diferentes tamanhos de tela.")
add_labeled(doc, "Severidade: ", "3 – problema grave, porque pode provocar erros recorrentes em ações importantes e irreversíveis, como o gasto de FitCoins.")

doc.add_heading("4.4 Crítica 4 – Iconografia, tipografia e apresentação dos dados", level=2)
add_labeled(doc, "Problema observado: ", "alguns dados técnicos aparecem diretamente na interface e certos elementos não se adaptam bem ao espaço disponível. No ranking, nomes como bot_atlas usam snake_case e avatares formados pelas duas primeiras letras geram a repetição “BO”. Na loja, títulos longos de multiplicadores são truncados, reduzindo a identificação do item.")
add_labeled(doc, "Evidência: ", "as capturas mostram iniciais repetidas no ranking e texto incompleto em cartões da loja. Isso enfraquece o reconhecimento visual e obriga o usuário a abrir ou deduzir informações que deveriam estar disponíveis no primeiro contato.")
add_labeled(doc, "Fundamentação: ", "as heurísticas de compatibilidade com o mundo real e consistência e padrões recomendam linguagem familiar e apresentação previsível. A WCAG 2.2 também orienta que a ampliação ou o estreitamento da área visível não provoque perda de informação em conteúdos que podem se reorganizar.")
add_labeled(doc, "Recomendação: ", "separar identificador interno de nome de exibição; apresentar “Bot Atlas” em vez de “bot_atlas”; usar avatares distintos ou ícones de bot; permitir duas linhas e altura adaptável nos títulos da loja; definir hierarquia tipográfica consistente; e testar a interface com fonte do sistema em 200%, telas estreitas e os temas claro e escuro.")
add_labeled(doc, "Severidade: ", "2 – problema pequeno, pois a função permanece disponível, mas a compreensão e o reconhecimento são prejudicados.")

doc.add_heading("4.5 Crítica 5 – Arquitetura da informação e oclusão de conteúdo", level=2)
add_labeled(doc, "Problema observado: ", "na tela inicial, o botão fixo “INICIAR ATIVIDADE” ocupa a região inferior do cartão mensal e pode encobrir informações ou dar a impressão de que parte do calendário não existe. A sequência vertical ainda mistura dados do dia, meta, histórico mensal e último percurso sem uma divisão progressiva clara.")
add_labeled(doc, "Evidência: ", "na captura, o botão vermelho se sobrepõe visualmente à porção inferior do calendário. O destaque intenso da ação é apropriado para a tarefa principal, mas a falta de espaço reservado compromete a leitura do conteúdo histórico e aumenta a necessidade de exploração por rolagem.")
add_labeled(doc, "Fundamentação: ", "as heurísticas de visibilidade do estado, design estético e minimalista e reconhecimento em vez de memorização recomendam que ações persistentes não ocultem informações relevantes. A teoria da carga cognitiva de Sweller sustenta a redução de demandas desnecessárias; os princípios de agrupamento da Gestalt favorecem separar dados diários, históricos e ações por proximidade e hierarquia.")
add_labeled(doc, "Recomendação: ", "reservar padding inferior equivalente à altura do botão e da navegação; posicionar a ação em uma área própria do Scaffold ou no fluxo do conteúdo; agrupar “Hoje”, “Meta semanal” e “Histórico” em seções nomeadas; e usar divulgação progressiva para detalhes do percurso. A correção deve ser verificada em telas pequenas e com fontes ampliadas.")
add_labeled(doc, "Severidade: ", "3 – problema grave, porque conteúdo importante pode ficar parcialmente oculto e parecer indisponível.")

doc.add_heading("5 QUADRO DE SEVERIDADE E RECOMENDAÇÕES", level=1)
headers = ["Nº", "Síntese do problema", "Eixos principais", "Sev.", "Prioridade de correção"]
rows = [
    ["1", "Texto secundário pequeno e com baixo contraste", "Cor, acessibilidade e tipografia", "3", "Alta"],
    ["2", "Bots sem identificação explícita no ranking", "Ética, cultura e psicologia", "2", "Média"],
    ["3", "Ações da loja com área de toque de 44 dp", "Ergonomia e métricas", "3", "Alta"],
    ["4", "Iniciais repetidas, snake_case e truncamento", "Iconografia e tipografia", "2", "Média"],
    ["5", "Botão persistente encobre o calendário", "Arquitetura e carga cognitiva", "3", "Alta"],
]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
table.autofit = False
widths = [Cm(0.8), Cm(6.0), Cm(4.0), Cm(1.1), Cm(3.0)]
for i, (cell, text, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
    cell.width = width
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "D9EAD3")
    set_cell_margins(cell)
    for r in cell.paragraphs[0].runs:
        r.bold = True
        r.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_repeat_table_header(table.rows[0])
for row_data in rows:
    row = table.add_row()
    for cell, text, width in zip(row.cells, row_data, widths):
        cell.width = width
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for r in cell.paragraphs[0].runs:
            r.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("6 CONSIDERAÇÕES FINAIS", level=1)
add_body(doc, "A avaliação identificou cinco problemas que atravessam os temas propostos pela disciplina. As prioridades imediatas são corrigir contraste e legibilidade, garantir alvos de toque adequados e impedir a oclusão do calendário. Em seguida, recomenda-se tornar o ranking transparente quanto aos perfis simulados e padronizar nomes, avatares e textos da loja.")
add_body(doc, "As conclusões não substituem uma avaliação com usuários. Após as correções, o próximo passo deve envolver tarefas representativas — iniciar atividade, consultar ranking, comprar um item e interpretar o perfil — com participantes do público-alvo. Podem ser coletadas taxa de conclusão, erros, tempo por tarefa, necessidade de ajuda e percepção subjetiva. Também é recomendável incluir pessoas que utilizem ampliação de fonte ou leitor de tela para verificar se as melhorias atendem diferentes condições de uso.")

doc.add_heading("REFERÊNCIAS", level=1)
references = [
    "ANDROID DEVELOPERS. Accessibility in Jetpack Compose. [S. l.]: Google, 2026. Disponível em: https://developer.android.com/develop/ui/compose/accessibility. Acesso em: 30 ago. 2026.",
    "ANDROID DEVELOPERS. Accessibility in Compose: API defaults. [S. l.]: Google, 2026. Disponível em: https://developer.android.com/develop/ui/compose/accessibility/api-defaults. Acesso em: 30 ago. 2026.",
    "FESTINGER, Leon. A theory of social comparison processes. Human Relations, v. 7, n. 2, p. 117–140, 1954. DOI: https://doi.org/10.1177/001872675400700202.",
    "FITTS, Paul M. The information capacity of the human motor system in controlling the amplitude of movement. Journal of Experimental Psychology, v. 47, n. 6, p. 381–391, 1954. DOI: https://doi.org/10.1037/h0055392.",
    "GRAY, Colin M. et al. The dark (patterns) side of UX design. In: CHI CONFERENCE ON HUMAN FACTORS IN COMPUTING SYSTEMS, 2018. Proceedings [...]. New York: ACM, 2018. p. 1–14. DOI: https://doi.org/10.1145/3173574.3174108.",
    "NIELSEN, Jakob. 10 usability heuristics for user interface design. Fremont: Nielsen Norman Group, 2024. Disponível em: https://www.nngroup.com/articles/ten-usability-heuristics/. Acesso em: 30 ago. 2026.",
    "NIELSEN, Jakob. Severity ratings for usability problems. Fremont: Nielsen Norman Group, 1994. Disponível em: https://www.nngroup.com/articles/how-to-rate-the-severity-of-usability-problems/. Acesso em: 30 ago. 2026.",
    "SWELLER, John. Cognitive load during problem solving: effects on learning. Cognitive Science, v. 12, n. 2, p. 257–285, 1988. DOI: https://doi.org/10.1207/s15516709cog1202_4.",
    "WORLD WIDE WEB CONSORTIUM. Web Content Accessibility Guidelines (WCAG) 2.2. W3C Recommendation, 5 out. 2023. Disponível em: https://www.w3.org/TR/WCAG22/. Acesso em: 30 ago. 2026.",
    "WORLD WIDE WEB CONSORTIUM. Understanding Success Criterion 1.4.3: Contrast (Minimum). [S. l.]: W3C, 2026. Disponível em: https://www.w3.org/WAI/WCAG22/Understanding/contrast-minimum. Acesso em: 30 ago. 2026.",
    "WORLD WIDE WEB CONSORTIUM. Understanding Success Criterion 1.4.10: Reflow. [S. l.]: W3C, 2026. Disponível em: https://www.w3.org/WAI/WCAG22/Understanding/reflow. Acesso em: 30 ago. 2026.",
]
for reference in references:
    add_reference(doc, reference)

# Rodapé com numeração apenas a partir da parte textual.
for sec in doc.sections:
    sec.footer_distance = Cm(1.0)
    sec.footer.is_linked_to_previous = False
    sec.footer.paragraphs[0].clear()
body_section.footer.is_linked_to_previous = False
add_page_number(body_section.footer.paragraphs[0])
restart_page_numbering(body_section, 1)

doc.core_properties.title = "Relatório de Avaliação de Usabilidade e Acessibilidade do FitWorkUp"
doc.core_properties.author = "Ronaldo Correia Couto"
doc.core_properties.subject = "Inspeção heurística preliminar da interface do aplicativo FitWorkUp"
doc.core_properties.keywords = "usabilidade, acessibilidade, FitWorkUp, IHC, WCAG"

doc.save(OUTPUT)
print(OUTPUT)
